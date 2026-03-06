
import pandas as pd
import json
import fitz
from docx import Document

def parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_csv(filepath):
    df = pd.read_csv(filepath)
    # scan column names themselves
    text = ' '.join(df.columns.tolist())
    # scan all cell values
    text += ' ' + df.to_string()
    return text

def parse_json(filepath):
    with open(filepath, 'r') as f:
        data = json.load(f)
    
    def extract_values(obj):
        if isinstance(obj, dict):
            return ' '.join([extract_values(v) for v in obj.values()])
        elif isinstance(obj, list):
            return ' '.join([extract_values(i) for i in obj])
        else:
            return str(obj)
    
    return extract_values(data)


# ---- TEST IT ----
if __name__ == "__main__":
    # test txt
    with open("sample_files/test.txt", "w") as f:
        f.write("My name is Rahul Sharma. Email: rahul@gmail.com")
    print("TXT result:")
    print(parse_txt("sample_files/test.txt"))
    print()

    # test csv
    with open("sample_files/test.csv", "w") as f:
        f.write("name,email,phone\nRahul Sharma,rahul@gmail.com,9876543210")
    print("CSV result:")
    print(parse_csv("sample_files/test.csv"))
    print()

    # test json
    with open("sample_files/test.json", "w") as f:
        json.dump({"name": "Rahul Sharma", "email": "rahul@gmail.com"}, f)
    print("JSON result:")
    print(parse_json("sample_files/test.json"))



    import fitz  # PyMuPDF
from docx import Document

def parse_pdf(filepath):
    text = ""
    doc = fitz.open(filepath)
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def parse_docx(filepath):
    doc = Document(filepath)
    text = ""
    # extract from paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    # extract from tables — most teams miss this
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    return text
    # add inside if __name__ == "__main__": block

    # test pdf
    print("PDF result:")
    # create a simple test - we will test with real pdf later
    print("PDF parser ready - waiting for real PDF file")
    print()

    # test docx
    print("DOCX result:")
    print("DOCX parser ready - waiting for real DOCX file")

def extract_text(filepath):
    """
    Master function — takes any file path
    returns plain text regardless of format
    """
    ext = filepath.split(".")[-1].lower()
    
    if ext == "txt":
        return parse_txt(filepath)
    elif ext == "csv":
        return parse_csv(filepath)
    elif ext == "json":
        return parse_json(filepath)
    elif ext == "pdf":
        return parse_pdf(filepath)
    elif ext == "docx":
        return parse_docx(filepath)
    else:
        # fallback — try reading as plain text
        try:
            return parse_txt(filepath)
        except:
            return ""


# ---- TEST IT ----
if __name__ == "__main__":
    # test txt
    with open("sample_files/test.txt", "w") as f:
        f.write("My name is Rahul Sharma. Email: rahul@gmail.com")
    print("TXT result:")
    print(parse_txt("sample_files/test.txt"))
    print()

    # test csv
    with open("sample_files/test.csv", "w") as f:
        f.write("name,email,phone\nRahul Sharma,rahul@gmail.com,9876543210")
    print("CSV result:")
    print(parse_csv("sample_files/test.csv"))
    print()

    # test json
    with open("sample_files/test.json", "w") as f:
        import json
        json.dump({"name": "Rahul Sharma", "email": "rahul@gmail.com"}, f)
    print("JSON result:")
    print(parse_json("sample_files/test.json"))
    print()

    print("✅ All parsers ready")
    print("✅ extract_text() master function ready")


